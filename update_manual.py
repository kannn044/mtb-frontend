#!/usr/bin/env python3
"""Update MTB_Cluster_User_Manual.docx with 10 new sections."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil, os

SRC = '/Users/gunkhemnak/Gun/mtb-frontend/MTB_Cluster_User_Manual.docx'
BAK = SRC.replace('.docx', '_backup.docx')
shutil.copy2(SRC, BAK)
print(f"Backup: {BAK}")

doc = Document(SRC)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# ==================== PAGE BREAK ====================
doc.add_page_break()

# ==================== SECTION HEADER ====================
add_heading(doc, 'ภาคผนวก: รายละเอียดทางเทคนิคของระบบ MTB Cluster Detection', level=1)
add_para(doc, 'เนื้อหาในส่วนนี้อธิบายรายละเอียดทางเทคนิคเพิ่มเติมของระบบ MTB Cluster Detection Web Application เพื่อให้ผู้ใช้งานและผู้ดูแลระบบเข้าใจการทำงานของระบบได้อย่างครบถ้วน โดยอ้างอิงจากซอร์สโค้ดและเอกสารขอบเขตงาน (TOR) ของโครงการ')

# ==================== 1. REMEMBER ME ====================
add_heading(doc, '1. ฟังก์ชัน "Remember Me" (จดจำการเข้าสู่ระบบ)', level=2)

add_para(doc, 'ฟังก์ชัน "Remember Me" คือตัวเลือกบนหน้าเข้าสู่ระบบ (Login) ที่ช่วยให้ผู้ใช้งานไม่ต้องกรอกชื่อผู้ใช้และรหัสผ่านซ้ำทุกครั้งที่เปิดเว็บไซต์ โดยระบบจะขยายระยะเวลาการใช้งาน (Session) ออกไปให้ยาวนานขึ้น')

add_heading(doc, 'หลักการทำงาน', level=3)
add_para(doc, 'เมื่อผู้ใช้งานเข้าสู่ระบบสำเร็จ ระบบจะสร้าง Token สำหรับยืนยันตัวตน (JWT — JSON Web Token) ซึ่งเปรียบเสมือน "บัตรผ่าน" ที่มีอายุการใช้งานจำกัด โดยมีความแตกต่างดังนี้:')

add_bullet(doc, ' Token จะมีอายุการใช้งาน 1 วัน (24 ชั่วโมง) และถูกเก็บใน Session Storage ของเบราว์เซอร์ ซึ่งจะถูกลบทันทีเมื่อปิดแท็บหรือปิดเบราว์เซอร์ ผู้ใช้งานจะต้องเข้าสู่ระบบใหม่ทุกครั้ง', bold_prefix='ไม่เลือก Remember Me: ')
add_bullet(doc, ' Token จะมีอายุการใช้งาน 30 วัน และถูกเก็บใน Local Storage ของเบราว์เซอร์ ซึ่งจะคงอยู่แม้ปิดเบราว์เซอร์แล้วเปิดใหม่ ทำให้ผู้ใช้งานสามารถเข้าถึงระบบได้ต่อเนื่องโดยไม่ต้องเข้าสู่ระบบซ้ำภายใน 30 วัน', bold_prefix='เลือก Remember Me: ')

add_heading(doc, 'คำแนะนำด้านความปลอดภัย', level=3)
add_bullet(doc, 'ไม่ควรเลือก "Remember Me" บนคอมพิวเตอร์สาธารณะหรือคอมพิวเตอร์ที่ใช้ร่วมกัน')
add_bullet(doc, 'หากสงสัยว่ามีผู้อื่นเข้าถึงบัญชีของท่าน ให้เปลี่ยนรหัสผ่านทันที ซึ่งจะทำให้ Token เดิมหมดอายุ')

# ==================== 2. ADMIN ควรมีคนเดียว ====================
add_heading(doc, '2. นโยบายผู้ดูแลระบบ (Administrator) ควรมีเพียงหนึ่งบัญชี', level=2)

add_para(doc, 'ระบบ MTB Cluster Detection ออกแบบให้มีบัญชีผู้ดูแลระบบ (ADMIN) เพียงบัญชีเดียว ด้วยเหตุผลด้านความปลอดภัย ดังนี้:')

add_heading(doc, '2.1 หลักการแบ่งแยกสิทธิ์ขั้นต่ำ (Principle of Least Privilege)', level=3)
add_para(doc, 'ตามหลัก Principle of Least Privilege (PoLP) ซึ่งเป็นหลักการพื้นฐานของ Information Security ผู้ใช้งานแต่ละคนควรได้รับสิทธิ์เท่าที่จำเป็นต่อการปฏิบัติงานเท่านั้น การมี ADMIN หลายคนจะเพิ่มความเสี่ยงที่สิทธิ์ระดับสูงจะถูกใช้งานอย่างไม่เหมาะสม ระบบจึงกำหนดบทบาท (Role) เป็น 3 ระดับ ได้แก่ ADMIN, UPLOADER และ VIEWER เพื่อให้แต่ละบทบาทมีสิทธิ์เฉพาะงานของตน')

add_heading(doc, '2.2 ความรับผิดชอบที่ตรวจสอบได้ (Single Point of Accountability)', level=3)
add_para(doc, 'เมื่อมี ADMIN เพียงคนเดียว การกระทำที่เกิดขึ้นในระดับ ADMIN (เช่น การสร้าง/ลบผู้ใช้ การเปลี่ยนสิทธิ์) สามารถตรวจสอบย้อนกลับได้อย่างชัดเจนว่าเป็นของบุคคลใด ผ่านระบบ Audit Log ที่บันทึกทุกกิจกรรม')

add_heading(doc, '2.3 ลดพื้นผิวการโจมตี (Reduce Attack Surface)', level=3)
add_para(doc, 'บัญชี ADMIN เป็นเป้าหมายหลักของการโจมตี (Attack Target) การมี ADMIN หลายบัญชีหมายถึงมีช่องทางการโจมตี (Attack Vector) มากขึ้น เช่น การเดารหัสผ่าน (Brute Force) หรือ Social Engineering ซึ่งหากบัญชี ADMIN ใดบัญชีหนึ่งถูกเจาะ ผู้โจมตีจะสามารถเข้าถึงและควบคุมระบบทั้งหมดได้')

add_heading(doc, '2.4 การออกแบบในระบบ', level=3)
add_para(doc, 'ในระบบ MTB Cluster Detection ข้อจำกัดนี้ถูกบังคับใช้ผ่านกลไกหลายชั้น:')
add_bullet(doc, 'บัญชี ADMIN จะไม่ปรากฏในรายการผู้ใช้งาน (User List) ของหน้าจัดการผู้ใช้ เพื่อป้องกันการแก้ไขหรือลบโดยไม่ตั้งใจ')
add_bullet(doc, 'ตัวเลือกบทบาท ADMIN ถูกปิดการใช้งานในหน้า User Management ทำให้ไม่สามารถสร้างบัญชี ADMIN ใหม่ผ่าน GUI ได้')
add_bullet(doc, 'Middleware ตรวจสอบสิทธิ์ ADMIN (checkAdmin) จะปฏิเสธคำร้องขอทันทีหากผู้ใช้ไม่มีสถานะ ADMIN ในระบบ')

# ==================== 3. SOFT DELETE ====================
add_heading(doc, '3. การลบข้อมูลผู้ใช้งาน (Soft Delete — การปิดใช้งานแทนการลบจริง)', level=2)

add_para(doc, 'เมื่อผู้ดูแลระบบทำการ "ลบ" ผู้ใช้งานในระบบ ข้อมูลของผู้ใช้คนนั้นจะไม่ถูกลบออกจากฐานข้อมูลจริง แต่ระบบจะเปลี่ยนสถานะ is_active จาก "Y" (ใช้งาน) เป็น "N" (ปิดใช้งาน) ซึ่งเรียกว่า Soft Delete')

add_heading(doc, 'เหตุผลที่ใช้ Soft Delete แทนการลบจริง (Hard Delete)', level=3)

add_bullet(doc, ' ข้อมูลประวัติการใช้งานของผู้ใช้ (เช่น ประวัติการ Upload, ประวัติการ Run Pipeline, Audit Log) ยังคงเชื่อมโยงกับ User ID ในฐานข้อมูล หากลบผู้ใช้จริง ข้อมูลเหล่านี้จะสูญหายหรือกลายเป็นข้อมูลกำพร้า (Orphan Data) ไม่สามารถตรวจสอบย้อนกลับได้', bold_prefix='1. รักษาความสมบูรณ์ของข้อมูล (Data Integrity): ')
add_bullet(doc, ' Audit Log ที่ผูกกับ user_id จะยังคงสามารถติดตามย้อนกลับไปยังผู้ใช้งานต้นทางได้ ซึ่งจำเป็นสำหรับการตรวจสอบความปลอดภัย (Security Audit)', bold_prefix='2. การตรวจสอบย้อนกลับ (Audit Trail): ')
add_bullet(doc, ' หากลบผู้ใช้ผิดคน สามารถเปิดใช้งานกลับได้ทันที (เปลี่ยน is_active กลับเป็น "Y") โดยไม่สูญเสียข้อมูลใดๆ', bold_prefix='3. ป้องกันการลบผิดพลาดและกู้คืนได้ (Recoverability): ')
add_bullet(doc, ' ผู้ใช้ที่ถูก Deactivate จะไม่สามารถเข้าสู่ระบบได้ แต่ผลลัพธ์การวิเคราะห์ (Pipeline Results) ที่เกี่ยวข้องจะยังคงอยู่ใน File System เพื่อให้ ADMIN สามารถตรวจสอบได้', bold_prefix='4. ข้อมูลผลลัพธ์การวิเคราะห์ไม่สูญหาย: ')

# ==================== 4. HASH PASSWORD ====================
add_heading(doc, '4. การเข้ารหัสรหัสผ่าน (Password Hashing)', level=2)

add_para(doc, 'ระบบ MTB Cluster Detection ไม่ได้จัดเก็บรหัสผ่านของผู้ใช้งานในรูปแบบข้อความธรรมดา (Plain Text) แต่จะแปลงรหัสผ่านให้อยู่ในรูปแบบ Hash ก่อนจัดเก็บลงฐานข้อมูล')

add_heading(doc, '4.1 Hash คืออะไร', level=3)
add_para(doc, 'Hash คือกระบวนการแปลงข้อมูล (ในที่นี้คือรหัสผ่าน) ให้เป็นค่าตัวอักษรขนาดคงที่ (Fixed-length String) ที่ไม่สามารถย้อนกลับเป็นรหัสผ่านเดิมได้ (One-Way Function) ตัวอย่างเช่น:')
add_bullet(doc, 'รหัสผ่าน "MyP@ss123" → Hash: "e99a18c428cb38d5f260853678922e03"')
add_bullet(doc, 'แม้จะรู้ค่า Hash ก็ไม่สามารถถอดรหัสกลับเป็น "MyP@ss123" ได้')

add_heading(doc, '4.2 ขั้นตอนการทำงาน', level=3)
add_bullet(doc, ' เมื่อผู้ใช้ลงทะเบียน ระบบจะ Hash รหัสผ่านแล้วจัดเก็บเฉพาะค่า Hash ลงฐานข้อมูล', bold_prefix='การลงทะเบียน: ')
add_bullet(doc, ' เมื่อผู้ใช้เข้าสู่ระบบ ระบบจะ Hash รหัสผ่านที่กรอกเข้ามา แล้วนำไปเปรียบเทียบกับค่า Hash ที่จัดเก็บไว้ หากตรงกันจะอนุญาตให้เข้าสู่ระบบ', bold_prefix='การเข้าสู่ระบบ: ')
add_bullet(doc, ' เมื่อเปลี่ยนรหัสผ่าน ระบบจะ Hash รหัสผ่านเก่าเพื่อตรวจสอบกับค่าในฐานข้อมูลก่อน จากนั้น Hash รหัสผ่านใหม่แล้วอัปเดต', bold_prefix='การเปลี่ยนรหัสผ่าน: ')

add_heading(doc, '4.3 ข้อกำหนดความซับซ้อนของรหัสผ่าน', level=3)
add_para(doc, 'ระบบบังคับให้รหัสผ่านใหม่ต้องมีความซับซ้อนขั้นต่ำ ดังนี้:')
add_bullet(doc, 'ความยาวอย่างน้อย 8 ตัวอักษร')
add_bullet(doc, 'มีตัวอักษรพิมพ์เล็ก (a-z) อย่างน้อย 1 ตัว')
add_bullet(doc, 'มีตัวอักษรพิมพ์ใหญ่ (A-Z) อย่างน้อย 1 ตัว')
add_bullet(doc, 'มีตัวเลข (0-9) อย่างน้อย 1 ตัว')
add_bullet(doc, 'มีอักขระพิเศษ (@, $, !, %, *, ?, &) อย่างน้อย 1 ตัว')
add_bullet(doc, 'รหัสผ่านใหม่ต้องไม่ซ้ำกับรหัสผ่านเดิม')

add_para(doc, 'ข้อกำหนดนี้สอดคล้องกับมาตรฐาน NIST Special Publication 800-63B (Digital Identity Guidelines) ที่แนะนำให้ใช้รหัสผ่านที่มีความซับซ้อนเพียงพอเพื่อป้องกันการโจมตีแบบ Brute Force และ Dictionary Attack')

# ==================== 5. QUEUE FIFO ====================
add_heading(doc, '5. ระบบจัดการคิวการประมวลผล (Pipeline Queue — FIFO)', level=2)

add_para(doc, 'ระบบ MTB Cluster Detection ใช้หลักการจัดคิวแบบ First In, First Out (FIFO) ในการจัดการลำดับการประมวลผล Pipeline ตามที่กำหนดในขอบเขตงาน (TOR) ข้อ 3.1')

add_heading(doc, '5.1 FIFO คืออะไร', level=3)
add_para(doc, 'FIFO (First In, First Out) คือหลักการจัดคิวที่ผู้ส่งคำร้องขอก่อนจะได้รับการประมวลผลก่อน เปรียบเสมือนการเข้าแถวซื้อของ คนที่มาก่อนจะได้รับบริการก่อน')

add_heading(doc, '5.2 หลักการทำงานในระบบ', level=3)
add_bullet(doc, 'เมื่อผู้ใช้กดปุ่ม "Run Process" ระบบจะบันทึกคำร้องขอลงตาราง pipeline_runs ด้วยสถานะ QUEUED พร้อมบันทึกเวลาที่ร้องขอ (requested_at)')
add_bullet(doc, 'ระบบ Worker จะตรวจสอบคิวทุก 5 วินาที (ค่าเริ่มต้น) หากไม่มี Pipeline ใดกำลังทำงานอยู่ (RUNNING) ระบบจะดึงคำร้องขอที่ requested_at เก่าที่สุดมาประมวลผล')
add_bullet(doc, 'ระบบอนุญาตให้ทำงานได้ทีละ 1 Pipeline เท่านั้น เพื่อป้องกันปัญหา Race Condition และการแย่งทรัพยากร CPU/RAM บน Processing Node')
add_bullet(doc, 'สถานะของแต่ละ Pipeline Run จะเปลี่ยนไปตามลำดับ: QUEUED → RUNNING → SUCCESS หรือ FAILED')
add_bullet(doc, 'เมื่อ Pipeline ทำงานเสร็จ ระบบจะส่งอีเมลแจ้งเตือนผู้ใช้งานที่ลงทะเบียนไว้โดยอัตโนมัติ')

# ==================== 6. DOWNLOAD TSV ====================
add_heading(doc, '6. การดาวน์โหลดไฟล์ผลลัพธ์ TSV จาก Zip File', level=2)

add_para(doc, 'ผู้ใช้งานสามารถดาวน์โหลดผลลัพธ์การวิเคราะห์ทั้งหมดจากหน้า Download ในรูปแบบไฟล์ Zip โดยข้อมูลภาพรวมและรายงานกลุ่มก้อนการระบาดในรูปแบบ TSV (Tab-Separated Values) จะอยู่ภายในไฟล์ Zip ดังนี้:')

add_heading(doc, 'โครงสร้างไฟล์ภายใน Zip', level=3)
add_para(doc, 'เมื่อแตกไฟล์ Zip แล้ว ผลลัพธ์ที่สำคัญจะอยู่ในโฟลเดอร์ user_reports/ โดยมีโครงสร้างดังนี้:')

add_bullet(doc, '0_input_data_validation_reports/ — รายงานการตรวจสอบข้อมูลนำเข้า')
add_bullet(doc, '1_seq_data_quality_reports/ — รายงานคุณภาพข้อมูลลำดับพันธุกรรม')
add_bullet(doc, '2_read_mapping_stats_reports/ — สถิติการ Map อ่านข้อมูลกับ Reference Genome')
add_bullet(doc, '3_lineage_genotyping_reports/ — รายงานสายพันธุ์ (Lineage)')
add_bullet(doc, '4_dr_profiling_reports/ — รายงานการดื้อยา (Drug Resistance Profile)')
add_bullet(doc, '5_enriched_input_metadata/ — ข้อมูล Metadata ที่เสริมผลวิเคราะห์แล้ว')
add_bullet(doc, '6_single_sample_summary_reports/ — รายงานสรุปรายตัวอย่าง')
add_bullet(doc, ' — รายงานกลุ่มก้อนการระบาด (Cluster Reports)', bold_prefix='7_WGS_cluster_reports/')

add_heading(doc, 'ไฟล์ cluster_membership.txt (TSV)', level=3)
add_para(doc, 'ไฟล์ cluster_membership.txt อยู่ในโฟลเดอร์ 7_WGS_cluster_reports/overall_report/ เป็นไฟล์รูปแบบ TSV ที่ระบุว่าตัวอย่างแต่ละตัวอย่างถูกจัดอยู่ในกลุ่มก้อนการระบาด (Cluster) ใดบ้าง สามารถเปิดด้วยโปรแกรม Microsoft Excel, Google Sheets หรือ Text Editor ทั่วไป')

add_para(doc, 'นอกจากนี้ ในโฟลเดอร์ 7_WGS_cluster_reports/overall_report/ ยังมีไฟล์อื่นๆ ที่เกี่ยวข้อง ได้แก่:')
add_bullet(doc, 'cluster_stats.txt — สถิติสรุปของแต่ละ Cluster')
add_bullet(doc, 'overall_wgs_cluster_summary_report.html — รายงาน HTML แบบอินเทอร์แอกทีฟ')
add_bullet(doc, 'overall_wgs_cluster_summary_report.tree.nwk — ไฟล์ Phylogenetic Tree ในรูปแบบ Newick')

# ==================== 7. USER VALIDATION ====================
add_heading(doc, '7. กฎเกณฑ์การตรวจสอบข้อมูลผู้ใช้งาน (User Data Validation)', level=2)

add_para(doc, 'ระบบมีกฎเกณฑ์การตรวจสอบข้อมูลผู้ใช้งานในขั้นตอนการลงทะเบียนและจัดการบัญชี ดังนี้:')

add_heading(doc, '7.1 ข้อมูลที่จำเป็น (Required Fields)', level=3)
add_bullet(doc, 'ชื่อผู้ใช้ (Username) — ต้องกรอก, ต้องไม่ซ้ำกับผู้ใช้อื่นในระบบ')
add_bullet(doc, 'อีเมล (Email) — ต้องกรอก, ต้องอยู่ในรูปแบบอีเมลที่ถูกต้อง, ต้องไม่ซ้ำกับผู้ใช้อื่น')
add_bullet(doc, 'ชื่อจริง (Name) — ต้องกรอก')
add_bullet(doc, 'นามสกุล (Lastname) — ต้องกรอก')
add_bullet(doc, 'หน่วยงาน (Organization) — ไม่บังคับ')

add_heading(doc, '7.2 การตรวจสอบอีเมล (Email Validation)', level=3)
add_bullet(doc, 'ต้องมีรูปแบบ user@domain.tld (ตรวจสอบด้วย Regular Expression)')
add_bullet(doc, 'ความยาวไม่เกิน 254 ตัวอักษร (ตามมาตรฐาน RFC 5321)')
add_bullet(doc, 'ไม่อนุญาตให้มีช่องว่าง (Whitespace) ในที่อยู่อีเมล')

add_heading(doc, '7.3 การตรวจสอบความซ้ำซ้อน (Uniqueness Check)', level=3)
add_bullet(doc, 'Username — ตรวจสอบว่าไม่ซ้ำกับ Username ที่มีอยู่แล้ว หากซ้ำจะแจ้งข้อผิดพลาด "Username is already taken"')
add_bullet(doc, 'Email — ตรวจสอบว่าไม่ซ้ำกับ Email ที่ลงทะเบียนแล้ว หากซ้ำจะแจ้งข้อผิดพลาด "Email is already registered"')

add_heading(doc, '7.4 การสร้างรหัสผ่านอัตโนมัติ', level=3)
add_para(doc, 'เมื่อผู้ใช้ลงทะเบียนผ่านหน้า Register ระบบจะสร้างรหัสผ่านแบบสุ่ม (Random Password) ขนาด 12 ตัวอักษรโดยอัตโนมัติ และส่งไปยังอีเมลที่ลงทะเบียนไว้ ผู้ใช้ควรเปลี่ยนรหัสผ่านหลังจากเข้าสู่ระบบครั้งแรก')

# ==================== 8. RECOMMENDED SERVER SPEC ====================
add_heading(doc, '8. โครงสร้างพื้นฐานที่เหมาะสม (Recommended Infrastructure Specification)', level=2)

add_para(doc, 'ตามขอบเขตงาน (TOR) ข้อ 1.3 และ 5 ระบบ MTB Cluster Detection ต้องการคอมพิวเตอร์อย่างน้อย 2 เครื่อง ทำงานร่วมกัน:')

add_heading(doc, '8.1 เครื่องที่ 1: User Terminal Node (เครื่องให้บริการผู้ใช้)', level=3)
p = doc.add_paragraph()
p.add_run('หน้าที่: ').bold = True
p.add_run('รับข้อมูลจากผู้ใช้ แสดงผลหน้าเว็บ และจัดการฐานข้อมูล')
add_bullet(doc, 'ระบบปฏิบัติการ: Ubuntu 22.04 LTS หรือ Ubuntu 24.04 LTS (64-bit)')
add_bullet(doc, 'CPU: 4 Cores ขึ้นไป (เช่น Intel Xeon E-2224 หรือ AMD EPYC 7232P)')
add_bullet(doc, 'RAM: 8 GB ขึ้นไป')
add_bullet(doc, 'Storage: SSD 100 GB ขึ้นไป (สำหรับ OS + Database + Application)')
add_bullet(doc, 'Network: Gigabit Ethernet')
add_bullet(doc, 'โปรแกรมที่ต้องติดตั้ง: Docker, Docker Compose, Node.js 18+, PM2, MySQL 8.0, Nginx (Reverse Proxy)')

add_heading(doc, '8.2 เครื่องที่ 2: Processing Node (เครื่องประมวลผล Genomics Pipeline)', level=3)
p = doc.add_paragraph()
p.add_run('หน้าที่: ').bold = True
p.add_run('ประมวลผล Nextflow Pipeline สำหรับวิเคราะห์ข้อมูล WGS')
add_bullet(doc, 'ระบบปฏิบัติการ: Ubuntu 22.04 LTS หรือ Ubuntu 24.04 LTS (64-bit)')
add_bullet(doc, 'CPU: 8 Cores ขึ้นไป (เช่น Intel Xeon E-2288G หรือ AMD EPYC 7302)')
add_bullet(doc, 'RAM: 16 GB ขึ้นไป (Pipeline กำหนดการใช้สูงสุดที่ 6 GB ต่อ Process แต่ระบบรวมต้องรองรับ OS + Conda Environment + Nextflow)')
add_bullet(doc, 'Storage: SSD 500 GB ขึ้นไป (สำหรับ OS + Genomic Reference Database + ข้อมูล FASTQ + ผลลัพธ์การวิเคราะห์)')
add_bullet(doc, 'Network: Gigabit Ethernet (เชื่อมต่อกับ User Terminal Node)')
add_bullet(doc, 'โปรแกรมที่ต้องติดตั้ง: Mamba/Conda, Java 17+ (สำหรับ Nextflow), Nextflow 22.10.1+, BWA, SAMtools, GATK, IQ-TREE, bcftools, R, Python 3')

add_heading(doc, '8.3 กระบวนการที่ต้อง Run', level=3)
add_bullet(doc, ' mtb-frontend (Next.js + PM2) — Port 3000', bold_prefix='User Terminal Node:')
add_bullet(doc, ' mtb-backend (Express.js + PM2) — Port 3001', bold_prefix='User Terminal Node:')
add_bullet(doc, ' MySQL 8.0 — Port 3306', bold_prefix='User Terminal Node:')
add_bullet(doc, ' Nextflow Pipeline (MTB_WGS_cluster_analysis) — ผ่าน Conda/Mamba Environment', bold_prefix='Processing Node:')

# ==================== 9. MINIMUM SPEC ====================
add_heading(doc, '9. สเปคขั้นต่ำที่ระบบทำงานได้ (Minimum Specification)', level=2)

add_para(doc, 'จากการวิเคราะห์ไฟล์ nextflow.config ของ Pipeline พบว่าทรัพยากรที่ Pipeline ต้องการมีดังนี้:')

add_heading(doc, 'ทรัพยากรที่ Pipeline กำหนด', level=3)
add_bullet(doc, 'process_low: CPU 4 Cores, RAM 4 GB')
add_bullet(doc, 'process_medium: CPU 4 Cores, RAM 6 GB')
add_bullet(doc, 'process_high: CPU 6 Cores, RAM 6 GB')
add_bullet(doc, 'Global Maximum: CPU 6 Cores, RAM 6 GB')

add_heading(doc, '9.1 สเปคขั้นต่ำ — Processing Node', level=3)
add_bullet(doc, 'CPU: 6 Cores (ขั้นต่ำสุดตาม process_high ที่กำหนดใน Pipeline)')
add_bullet(doc, 'RAM: 8 GB (Pipeline ใช้สูงสุด 6 GB + OS/Conda/Nextflow อีก ~2 GB)')
add_bullet(doc, 'Storage: 200 GB SSD (สำหรับ Reference Database + ข้อมูลผู้ใช้)')
add_bullet(doc, 'OS: Ubuntu 20.04 LTS ขึ้นไป')

add_heading(doc, '9.2 สเปคขั้นต่ำ — User Terminal Node', level=3)
add_bullet(doc, 'CPU: 2 Cores')
add_bullet(doc, 'RAM: 4 GB')
add_bullet(doc, 'Storage: 50 GB SSD')
add_bullet(doc, 'OS: Ubuntu 20.04 LTS ขึ้นไป หรือ ระบบปฏิบัติการที่รองรับ Docker')

add_heading(doc, '9.3 หมายเหตุ', level=3)
add_para(doc, 'สเปคขั้นต่ำข้างต้นเป็นค่าที่ระบบสามารถทำงานได้ แต่อาจใช้เวลานานในการประมวลผลเมื่อมีข้อมูลจำนวนมาก สำหรับการใช้งานจริงในระดับ Production แนะนำให้ใช้สเปคตามข้อ 8 (Recommended Specification) เพื่อประสิทธิภาพที่เหมาะสม')

# ==================== 10. SECURITY MEASURES ====================
add_heading(doc, '10. มาตรการรักษาความปลอดภัยของระบบ (Security Measures)', level=2)

add_para(doc, 'ระบบ MTB Cluster Detection ได้ออกแบบและพัฒนามาตรการรักษาความปลอดภัยหลายชั้น (Defense in Depth) ตามหลักทฤษฎีความปลอดภัยสารสนเทศระดับสากล ดังนี้:')

add_heading(doc, '10.1 การยืนยันตัวตน (Authentication) — อ้างอิง OWASP A07:2021', level=3)
add_bullet(doc, ' ระบบใช้ JSON Web Token (JWT) ซึ่งเป็นมาตรฐาน RFC 7519 ในการยืนยันตัวตนผู้ใช้งาน Token จะถูกส่งไปพร้อมกับทุกคำร้องขอ (Request) ผ่าน Authorization Header เพื่อตรวจสอบสิทธิ์ก่อนเข้าถึงทรัพยากร', bold_prefix='JWT Authentication: ')
add_bullet(doc, ' รหัสผ่านถูกจัดเก็บในรูปแบบ One-Way Hash ไม่สามารถถอดรหัสกลับได้ แม้ฐานข้อมูลจะถูกเจาะ ผู้โจมตีก็ไม่สามารถรู้รหัสผ่านจริงของผู้ใช้', bold_prefix='Password Hashing: ')

add_heading(doc, '10.2 การล็อกบัญชี (Account Lockout) — อ้างอิง NIST SP 800-63B Section 5.2.2', level=3)
add_para(doc, 'เพื่อป้องกันการโจมตีแบบ Brute Force (การเดารหัสผ่าน) ระบบมีมาตรการดังนี้:')
add_bullet(doc, 'หากกรอกรหัสผ่านผิดติดต่อกัน 5 ครั้ง บัญชีจะถูกล็อกอัตโนมัติเป็นเวลา 15 นาที')
add_bullet(doc, 'ระหว่างที่บัญชีถูกล็อก จะไม่สามารถเข้าสู่ระบบได้ แม้จะกรอกรหัสผ่านถูกต้อง')
add_bullet(doc, 'จำนวนครั้งที่กรอกผิดจะถูกรีเซ็ตเมื่อเข้าสู่ระบบสำเร็จ')
add_bullet(doc, 'ทุกความพยายามเข้าสู่ระบบ (สำเร็จ/ไม่สำเร็จ) ถูกบันทึกใน Audit Log พร้อม IP Address')

add_heading(doc, '10.3 การควบคุมสิทธิ์การเข้าถึง (Authorization — RBAC) — อ้างอิง OWASP A01:2021', level=3)
add_para(doc, 'ระบบใช้ Role-Based Access Control (RBAC) ตามหลักการ Principle of Least Privilege โดยแบ่งบทบาทเป็น 3 ระดับ:')
add_bullet(doc, ' มีสิทธิ์สูงสุด จัดการผู้ใช้ทั้งหมด ไม่ปรากฏใน User List เพื่อความปลอดภัย', bold_prefix='ADMIN: ')
add_bullet(doc, ' สามารถอัปโหลดข้อมูลและสั่ง Run Pipeline ได้', bold_prefix='UPLOADER: ')
add_bullet(doc, ' สามารถดู Dashboard และดาวน์โหลดผลลัพธ์ได้เท่านั้น', bold_prefix='VIEWER: ')

add_heading(doc, '10.4 บันทึกการตรวจสอบ (Audit Logging) — อ้างอิง CIA Triad & ISO 27001 A.12.4', level=3)
add_para(doc, 'ระบบบันทึกกิจกรรมสำคัญทั้งหมดลงตาราง audit_logs ในฐานข้อมูล เพื่อรักษาคุณสมบัติ Accountability ตามหลัก CIA Triad (Confidentiality, Integrity, Availability) โดยข้อมูลที่บันทึกประกอบด้วย:')
add_bullet(doc, 'User ID ของผู้กระทำ')
add_bullet(doc, 'ประเภทกิจกรรม (Action): LOGIN_SUCCESS, LOGIN_FAILED, ACCOUNT_LOCKED, REGISTER_SUCCESS เป็นต้น')
add_bullet(doc, 'รายละเอียด (Details): ข้อมูลเพิ่มเติมของเหตุการณ์')
add_bullet(doc, 'IP Address ของผู้กระทำ')
add_bullet(doc, 'Timestamp ของเหตุการณ์')

add_heading(doc, '10.5 การป้องกัน Path Traversal — อ้างอิง OWASP A01:2021 (Broken Access Control)', level=3)
add_para(doc, 'ระบบมีการป้องกัน Path Traversal Attack ในส่วนการดาวน์โหลดไฟล์ โดย:')
add_bullet(doc, 'ตรวจสอบ Run ID ด้วย Regular Expression เพื่อป้องกันอักขระพิเศษ')
add_bullet(doc, 'ตรวจสอบว่า File Path ที่ร้องขออยู่ภายในขอบเขตที่อนุญาตเท่านั้น (isPathInside check)')
add_bullet(doc, 'กำหนด Security Headers เช่น X-Content-Type-Options: nosniff และ Cache-Control: no-store')

add_heading(doc, '10.6 ความปลอดภัยของอีเมล (Email Security)', level=3)
add_para(doc, 'ระบบส่งอีเมลผ่านโปรโตคอล SMTP ที่เข้ารหัสด้วย TLS 1.2 ขึ้นไป (SMTP over TLS) เพื่อป้องกันการดักจับข้อมูลระหว่างทาง (Man-in-the-Middle Attack)')

add_heading(doc, '10.7 ตารางสรุปมาตรการรักษาความปลอดภัย', level=3)

# Add summary table
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'

headers = ['มาตรการ', 'เทคโนโลยี/วิธีการ', 'อ้างอิงทฤษฎี']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

rows_data = [
    ['การยืนยันตัวตน', 'JWT (JSON Web Token)', 'RFC 7519, OWASP A07:2021'],
    ['การเข้ารหัสรหัสผ่าน', 'One-Way Hash (MD5)', 'NIST SP 800-63B'],
    ['การล็อกบัญชี', 'Failed Attempts Counter + Time Lock', 'NIST SP 800-63B Section 5.2.2'],
    ['การควบคุมสิทธิ์', 'RBAC (Admin/Uploader/Viewer)', 'OWASP A01:2021, PoLP'],
    ['บันทึกตรวจสอบ', 'Audit Log (DB Table)', 'ISO 27001 A.12.4, CIA Triad'],
    ['ป้องกัน Path Traversal', 'Path Validation + Regex', 'OWASP A01:2021'],
    ['ความปลอดภัยอีเมล', 'SMTP over TLS 1.2+', 'RFC 8314'],
]

for row_idx, row_data in enumerate(rows_data, start=1):
    for col_idx, cell_text in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = cell_text

# Save the document
doc.save(SRC)
print(f"Updated: {SRC}")
print("Done! Manual has been updated with 10 new sections.")
